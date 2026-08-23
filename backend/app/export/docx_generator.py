import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(
    resume_data: dict,
    output_path: str,
    template_id: str = "ats_focused",
    sections_order: list = None,
    included_sections: list = None
) -> dict:
    """
    Generates a professional DOCX resume matching the finalized content and template style.
    """
    if sections_order is None:
        sections_order = ['summary', 'skills', 'experience', 'projects', 'education', 'certifications', 'achievements']
    if included_sections is None:
        included_sections = ['summary', 'skills', 'experience', 'projects', 'education', 'certifications', 'achievements']

    doc = Document()
    
    # Set Margins (0.5 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Style colors
    if template_id == 'classic_professional':
        font_name = 'Times New Roman'
        heading_rgb = RGBColor(30, 64, 175)
    elif template_id == 'modern_minimal':
        font_name = 'Arial'
        heading_rgb = RGBColor(13, 148, 136)
    elif template_id == 'technical':
        font_name = 'Courier New'
        heading_rgb = RGBColor(79, 70, 229)
    else: # ats_focused
        font_name = 'Arial'
        heading_rgb = RGBColor(37, 99, 235)

    # 1. Contact Header
    contact = resume_data.get('contact', {})
    name = contact.get('name') or "Atharva Raviraj Raut"
    email = contact.get('email') or ""
    phone = contact.get('phone') or ""
    linkedin = contact.get('linkedin') or ""
    github = contact.get('github') or ""

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(name.upper())
    run_title.font.name = font_name
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(17, 24, 39)

    contact_parts = [p for p in [phone, email, linkedin, github] if p]
    if contact_parts:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run("  •  ".join(contact_parts))
        run_sub.font.name = font_name
        run_sub.font.size = Pt(9)
        run_sub.font.color.rgb = RGBColor(75, 85, 99)

    def add_heading(text):
        h = doc.add_paragraph()
        run = h.add_run(text.upper())
        run.font.name = font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = heading_rgb

    def render_section(section_key):
        if section_key not in included_sections:
            return

        if section_key == 'summary':
            summary_text = resume_data.get('summary')
            if summary_text:
                add_heading("Professional Summary")
                p = doc.add_paragraph()
                run = p.add_run(summary_text)
                run.font.name = font_name
                run.font.size = Pt(9.5)

        elif section_key == 'skills':
            skills_raw = resume_data.get('skills_raw') or resume_data.get('skills')
            if skills_raw:
                add_heading("Technical Skills")
                if isinstance(skills_raw, list):
                    skills_str = ", ".join([s.get('canonical_name', s.get('skill_name', str(s))) for s in skills_raw])
                else:
                    skills_str = str(skills_raw)
                p = doc.add_paragraph()
                run = p.add_run(skills_str)
                run.font.name = font_name
                run.font.size = Pt(9.5)

        elif section_key == 'experience':
            exp_list = resume_data.get('experience', [])
            if exp_list:
                add_heading("Work Experience")
                for exp in exp_list:
                    exp_text = exp.get('raw') or exp.get('description') or ""
                    if exp_text:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(exp_text.strip())
                        run.font.name = font_name
                        run.font.size = Pt(9.5)

        elif section_key == 'projects':
            proj_list = resume_data.get('projects', [])
            if proj_list:
                add_heading("Projects")
                for proj in proj_list:
                    p_name = proj.get('name') or "Key Project"
                    p_desc = proj.get('raw') or proj.get('description') or ""
                    p_title = doc.add_paragraph()
                    r_title = p_title.add_run(p_name)
                    r_title.bold = True
                    r_title.font.name = font_name
                    r_title.font.size = Pt(10)
                    
                    if p_desc:
                        bullets = [b.strip() for b in p_desc.split('\n') if b.strip()]
                        for b in bullets:
                            clean_b = b.lstrip('•- ').strip()
                            bp = doc.add_paragraph(style='List Bullet')
                            r_b = bp.add_run(clean_b)
                            r_b.font.name = font_name
                            r_b.font.size = Pt(9.5)

        elif section_key == 'education':
            edu_list = resume_data.get('education', [])
            if edu_list:
                add_heading("Education")
                for edu in edu_list:
                    e_text = edu.get('raw') or f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                    if e_text:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(e_text.strip())
                        run.font.name = font_name
                        run.font.size = Pt(9.5)

        elif section_key == 'certifications':
            cert_list = resume_data.get('certifications', [])
            if cert_list:
                add_heading("Certifications")
                for c in cert_list:
                    c_text = c.get('raw') or c.get('name') or str(c)
                    if c_text:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(c_text.strip())
                        run.font.name = font_name
                        run.font.size = Pt(9.5)

        elif section_key == 'achievements':
            ach_list = resume_data.get('achievements', [])
            if ach_list:
                add_heading("Achievements")
                for a in ach_list:
                    a_text = a.get('raw') or str(a)
                    if a_text:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(a_text.strip())
                        run.font.name = font_name
                        run.font.size = Pt(9.5)

    for sec in sections_order:
        render_section(sec)

    doc.save(output_path)
    return {"output_path": output_path}
